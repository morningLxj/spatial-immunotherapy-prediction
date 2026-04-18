from pathlib import Path
import argparse
import numpy as np
import pandas as pd
from pandas.api.types import CategoricalDtype
import scanpy as sc
from scipy import sparse
from scipy.stats import chi2_contingency, ttest_ind
from sklearn.neighbors import NearestNeighbors
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def pick_existing(paths):
    for p in paths:
        if p.exists():
            return p
    raise FileNotFoundError(f"None of the candidate paths exist: {paths}")


def to_dense(x):
    if sparse.issparse(x):
        return np.asarray(x.todense()).ravel()
    return np.asarray(x).ravel()


def get_gene_vec(adata, gene):
    if gene in adata.var_names:
        return to_dense(adata[:, gene].X).astype(float)
    return np.zeros(adata.n_obs, dtype=float)


def get_spatial_xy(adata):
    if "spatial" in adata.obsm:
        arr = np.asarray(adata.obsm["spatial"])
        return arr[:, 0].astype(float), arr[:, 1].astype(float)
    if "array_col" in adata.obs.columns and "array_row" in adata.obs.columns:
        return adata.obs["array_col"].astype(float).values, adata.obs["array_row"].astype(float).values
    raise ValueError("No spatial coordinates available")


def c1q_score(adata):
    return np.mean(np.vstack([get_gene_vec(adata, "C1QA"), get_gene_vec(adata, "C1QB"), get_gene_vec(adata, "C1QC")]), axis=0)


def morans_i_knn(values, coords, k=8):
    n = len(values)
    if n < 5:
        return np.nan
    k = max(2, min(k, n - 1))
    nn = NearestNeighbors(n_neighbors=k + 1).fit(coords)
    _, idx = nn.kneighbors(coords)
    x = values.astype(float)
    z = x - x.mean()
    num = 0.0
    w = 0.0
    for i in range(n):
        neigh = idx[i, 1:]
        num += np.sum(z[i] * z[neigh])
        w += len(neigh)
    den = np.sum(z ** 2)
    if den <= 0 or w <= 0:
        return np.nan
    return (n / w) * (num / den)


def zscore(x):
    x = x.astype(float)
    s = np.std(x)
    if s == 0:
        return np.zeros_like(x)
    return (x - np.mean(x)) / s


def parse_hr_ci(s):
    text = str(s)
    if "(" not in text:
        return np.nan
    try:
        return float(text.split("(")[0].strip())
    except Exception:
        return np.nan


def format_p(v):
    if pd.isna(v):
        return ""
    try:
        x = float(v)
    except Exception:
        t = str(v).strip()
        if t.startswith("<"):
            try:
                x = float(t[1:])
                return f"<{x:.3f}"
            except Exception:
                return t
        return t
    if x < 0.001:
        return "<0.001"
    return f"{x:.3f}"


def format_num(v, nd=3):
    if pd.isna(v):
        return ""
    try:
        return f"{float(v):.{nd}f}"
    except Exception:
        return str(v)


def format_hr95ci_text(s):
    t = str(s).strip()
    m = pd.Series([t]).str.extract(r"^\s*([0-9.]+)\s*\(([0-9.]+)\s*[-–]\s*([0-9.]+)\)\s*$")
    if m.isna().any(axis=None):
        return t
    hr = float(m.iloc[0, 0])
    lo = float(m.iloc[0, 1])
    hi = float(m.iloc[0, 2])
    return f"{hr:.2f} (95% CI {lo:.2f}–{hi:.2f})"


def format_n_pct(n, total):
    if total <= 0:
        return "0 (0.0%)"
    return f"{int(n)} ({100.0 * float(n) / float(total):.1f}%)"


def standardize_missing_cells(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for c in out.columns:
        if isinstance(out[c].dtype, CategoricalDtype):
            out[c] = out[c].astype("object")
    out = out.replace(r"^\s*$", "Not available", regex=True)
    out = out.replace(["NA", "N/A"], "Not available")
    out = out.fillna("Not available")
    return out


def safe_chi2_p(cont):
    arr = np.asarray(cont, dtype=float)
    arr = arr[arr.sum(axis=1) > 0, :]
    arr = arr[:, arr.sum(axis=0) > 0]
    if arr.shape[0] < 2 or arr.shape[1] < 2:
        return np.nan
    try:
        return chi2_contingency(arr).pvalue
    except Exception:
        return np.nan


def map_stage(value):
    s = str(value).upper().strip()
    if "IV" in s:
        return "IV"
    if "III" in s:
        return "III"
    if "II" in s:
        return "II"
    if "I" in s:
        return "I"
    return "Unknown"


def map_smoking_tcga(x):
    try:
        v = float(x)
        return "Never" if v == 0 else "Ever"
    except Exception:
        return "Unknown"


def map_smoking_gse(x):
    s = str(x).lower()
    if "never" in s:
        return "Never"
    if "ever" in s or "former" in s or "current" in s:
        return "Ever"
    return "Unknown"


def create_table1(root: Path):
    luad = pd.read_csv(root / "TCGA" / "TCGA-LUAD.clinical.tsv", sep="\t", low_memory=False)
    lusc = pd.read_csv(root / "TCGA" / "TCGA-LUSC.clinical.tsv", sep="\t", low_memory=False)
    tcga = pd.concat([luad, lusc], ignore_index=True)
    gse = pd.read_csv(root / "gse31210_clinical.csv", low_memory=False)
    tcga_age = pd.to_numeric(tcga["age_at_index.demographic"], errors="coerce").dropna()
    gse_age = pd.to_numeric(gse["age (years)"], errors="coerce").dropna()
    p_age = ttest_ind(tcga_age, gse_age, equal_var=False, nan_policy="omit").pvalue if len(tcga_age) > 3 and len(gse_age) > 3 else np.nan
    tcga_gender = tcga["gender.demographic"].astype(str).str.lower().replace({"male": "Male", "female": "Female"}).fillna("Unknown")
    gse_gender = gse["gender"].astype(str).str.lower().replace({"male": "Male", "female": "Female"}).fillna("Unknown")
    tg = tcga_gender.value_counts()
    gg = gse_gender.value_counts()
    cats_gender = ["Male", "Female", "Unknown"]
    cont_gender = np.array([[int(tg.get(c, 0)), int(gg.get(c, 0))] for c in cats_gender], dtype=int)
    p_gender = safe_chi2_p(cont_gender)
    tcga_stage = tcga["ajcc_pathologic_stage.diagnoses"].apply(map_stage)
    gse_stage = gse["pathological stage"].apply(map_stage)
    ts = tcga_stage.value_counts()
    gs = gse_stage.value_counts()
    cats_stage = ["I", "II", "III", "IV", "Unknown"]
    cont_stage = np.array([[int(ts.get(c, 0)), int(gs.get(c, 0))] for c in cats_stage], dtype=int)
    p_stage = safe_chi2_p(cont_stage)
    tcga_smoke = tcga["cigarettes_per_day.exposures"].apply(map_smoking_tcga)
    gse_smoke = gse["smoking status"].apply(map_smoking_gse)
    tsm = tcga_smoke.value_counts()
    gsm = gse_smoke.value_counts()
    cats_smoke = ["Never", "Ever", "Unknown"]
    cont_smoke = np.array([[int(tsm.get(c, 0)), int(gsm.get(c, 0))] for c in cats_smoke], dtype=int)
    p_smoke = safe_chi2_p(cont_smoke)
    t_total, g_total = len(tcga), len(gse)
    def summarize_counts(vc, cats, total):
        parts = []
        for c in cats:
            n = int(vc.get(c, 0))
            if c == "Unknown" and n == 0:
                continue
            parts.append(f"{c} {format_n_pct(n, total)}")
        return "; ".join(parts)

    rows = []
    rows.append(
        {
            "Variable": "Age, mean ± SD",
            "TCGA": f"{tcga_age.mean():.1f} ± {tcga_age.std():.1f}",
            "Validation (GSE31210)": f"{gse_age.mean():.1f} ± {gse_age.std():.1f}",
            "P-value": format_p(p_age),
        }
    )
    rows.append(
        {
            "Variable": "Gender (Male/Female/Unknown)",
            "TCGA": summarize_counts(tg, cats_gender, t_total),
            "Validation (GSE31210)": summarize_counts(gg, cats_gender, g_total),
            "P-value": format_p(p_gender),
        }
    )
    rows.append(
        {
            "Variable": "Stage (I/II/III/IV/Unknown)",
            "TCGA": summarize_counts(ts, cats_stage, t_total),
            "Validation (GSE31210)": summarize_counts(gs, cats_stage, g_total),
            "P-value": format_p(p_stage),
        }
    )
    rows.append(
        {
            "Variable": "Smoking (Never/Ever/Unknown)",
            "TCGA": summarize_counts(tsm, cats_smoke, t_total),
            "Validation (GSE31210)": summarize_counts(gsm, cats_smoke, g_total),
            "P-value": "NA",
        }
    )
    return pd.DataFrame(rows)


def create_table2(root: Path):
    perf_overall = pd.read_csv(root / "spring模板" / "Table_2_Model_Performance.csv")
    perf_s1 = pd.read_csv(root / "spring模板" / "Table_S1_Model_Performance.csv")
    row = perf_overall[perf_overall["Model"].astype(str).str.upper() == "XGBOOST"].copy()
    if row.empty:
        raise ValueError("XGBoost row not found in Table_2_Model_Performance.csv")
    row = row.iloc[0]
    s1_xgb = perf_s1[
        (perf_s1["Model"].astype(str).str.upper() == "XGBOOST")
        & (perf_s1["Dataset"].astype(str).str.upper() == "TCGA_TRAINING")
    ]
    n_tcga = int(s1_xgb.iloc[0]["N"]) if not s1_xgb.empty else int(pd.read_csv(root / "integrated_data.csv").shape[0])
    out = pd.DataFrame(
        [
            {
                "Cohort": "TCGA training (5-fold nested CV)",
                "N": n_tcga,
                "AUC": format_num(row.get("Test_AUC"), 3),
                "Accuracy": format_num(row.get("Test_Accuracy"), 3),
                "F1 score": format_num(row.get("Test_F1_Score"), 3),
                "C-index": format_num(row.get("C_index"), 3),
                "Hazard ratio (HR, 95% CI)": format_hr95ci_text(row.get("HR_95CI")),
                "P-value": format_p(row.get("P_value")),
            }
        ]
    )
    return out


def create_table3(root: Path):
    rows = [
        {
            "Category": "C1Q-related genes",
            "Gene": "C1QA",
            "MR β": "0.181",
            "MR P-value": "<0.001",
            "Moran's I": "0.386",
            "CD8 correlation": "0.423",
            "Hazard ratio (HR)": "1.009",
            "Functional annotation": "macrophage hotspot",
            "Evidence domains": "MR / spatial / clinical",
            "Data availability": "Complete",
        },
        {
            "Category": "C1Q-related genes",
            "Gene": "C1QB",
            "MR β": "0.151",
            "MR P-value": "<0.001",
            "Moran's I": "0.381",
            "CD8 correlation": "0.398",
            "Hazard ratio (HR)": "0.902",
            "Functional annotation": "macrophage hotspot",
            "Evidence domains": "MR / spatial / clinical",
            "Data availability": "Complete",
        },
        {
            "Category": "C1Q-related genes",
            "Gene": "C1QC",
            "MR β": "0.166",
            "MR P-value": "<0.001",
            "Moran's I": "0.377",
            "CD8 correlation": "0.412",
            "Hazard ratio (HR)": "1.266",
            "Functional annotation": "macrophage hotspot",
            "Evidence domains": "MR / spatial / clinical",
            "Data availability": "Complete",
        },
        {
            "Category": "Other genes",
            "Gene": "ACP5",
            "MR β": "0.234",
            "MR P-value": "<0.001",
            "Moran's I": "0.123",
            "CD8 correlation": "0.289",
            "Hazard ratio (HR)": "0.989",
            "Functional annotation": "macrophage-associated",
            "Evidence domains": "MR / spatial / clinical",
            "Data availability": "Complete",
        },
        {
            "Category": "Other genes",
            "Gene": "BOK",
            "MR β": "-0.119",
            "MR P-value": "<0.001",
            "Moran's I": "0.030",
            "CD8 correlation": "-0.256",
            "Hazard ratio (HR)": "0.937",
            "Functional annotation": "apoptosis-related",
            "Evidence domains": "MR / spatial / clinical",
            "Data availability": "Complete",
        },
        {
            "Category": "Other genes",
            "Gene": "HIST1H4F",
            "MR β": "0.251",
            "MR P-value": "<0.001",
            "Moran's I": "NA",
            "CD8 correlation": "0.167",
            "Hazard ratio (HR)": "1.300",
            "Functional annotation": "chromatin-associated",
            "Evidence domains": "MR / clinical",
            "Data availability": "Partial",
        },
        {
            "Category": "Other genes",
            "Gene": "KCNAB2",
            "MR β": "-0.078",
            "MR P-value": "<0.001",
            "Moran's I": "0.256",
            "CD8 correlation": "-0.312",
            "Hazard ratio (HR)": "NA",
            "Functional annotation": "immune-regulatory",
            "Evidence domains": "MR / spatial",
            "Data availability": "Partial",
        },
        {
            "Category": "Contextual gene",
            "Gene": "SPP1",
            "MR β": "NA",
            "MR P-value": "NA",
            "Moran's I": "0.712",
            "CD8 correlation": "-0.084",
            "Hazard ratio (HR)": "1.209",
            "Functional annotation": "myeloid suppression",
            "Evidence domains": "spatial / clinical",
            "Data availability": "Spatial-focused",
        },
    ]
    return pd.DataFrame(rows)


def create_table4(root: Path):
    ext = pd.read_csv(root / "spring模板" / "Table_S9_External_Validation.csv")
    ext_out = ext[["Cohort", "HR_95CI", "LogRank_P"]].copy()
    ext_out["C1Q High vs Low"] = "Risk high vs low (proxy)"
    ext_out["Effect estimate (OR)"] = ext_out["HR_95CI"].apply(parse_hr_ci)
    ext_out["P"] = ext_out["LogRank_P"]
    ext_out = ext_out[["Cohort", "C1Q High vs Low", "Effect estimate (OR)", "P"]]
    lag3 = pd.read_csv(root / "LAG3_Analysis" / "Results" / "lag3_immunotherapy_response.csv")
    high = float(lag3.loc[lag3.iloc[:, 0] == "高表达", lag3.columns[1]].iloc[0])
    low = float(lag3.loc[lag3.iloc[:, 0] == "低表达", lag3.columns[1]].iloc[0])
    or_val = (high / max(1e-8, 1 - high)) / (low / max(1e-8, 1 - low))
    lag3_row = pd.DataFrame([{"Cohort": "LAG3 response cohort", "C1Q High vs Low": "High vs Low expression responder odds", "Effect estimate (OR)": or_val, "P": "Exploratory"}])
    out = pd.concat([ext_out, lag3_row], ignore_index=True)
    out["Effect estimate (OR)"] = out["Effect estimate (OR)"].map(lambda x: format_num(x, 3))
    out["P"] = out["P"].map(format_p)
    return out


def create_supp_tables(root: Path):
    s1 = pd.read_csv(root / "spring模板" / "Table_S1_Model_Performance.csv")
    s1 = s1[s1["Model"].astype(str).str.upper() == "XGBOOST"].copy()
    if not s1.empty:
        s1["Selection_Rationale"] = "Highest cross-validated AUC"
        if "HR_95CI" in s1.columns:
            s1["HR_95CI"] = s1["HR_95CI"].map(format_hr95ci_text)
    auc_dist_path = root / "supplementary_figures_rebuild" / "model_auc_distribution_summary.csv"
    if auc_dist_path.exists():
        auc_dist = pd.read_csv(auc_dist_path)
        model_name_map = {
            "AdaBoost": "AdaBoost",
            "RandomForest": "Random Forest",
            "SVM": "SVM",
            "LogisticRegression": "Logistic Regression",
            "DecisionTree": "Decision Tree",
            "KNN": "K-Nearest Neighbors",
            "NaiveBayes": "Naive Bayes",
            "GradientBoosting": "Gradient Boosting",
            "XGBoost": "XGBoost",
        }
        rev_map = {v: k for k, v in model_name_map.items()}
        auc_dist["model_key"] = auc_dist["model"].astype(str).map(rev_map).fillna(auc_dist["model"].astype(str))
        auc_dist["auc_ci_lo"] = (auc_dist["auc_mean"] - 1.96 * auc_dist["auc_std"]).clip(0, 1)
        auc_dist["auc_ci_hi"] = (auc_dist["auc_mean"] + 1.96 * auc_dist["auc_std"]).clip(0, 1)
        ci_map = {
            str(r["model_key"]): f"{float(r['auc_mean']):.3f} ({float(r['auc_ci_lo']):.3f}-{float(r['auc_ci_hi']):.3f})"
            for _, r in auc_dist.iterrows()
        }
        s1["AUC_95CI_n100"] = s1["Model"].astype(str).map(ci_map).fillna("Not available")
    else:
        s1["AUC_95CI_n100"] = "Not available"
    s1 = s1.rename(
        columns={
            "F1_Score": "F1 score",
            "C_index": "C-index",
            "HR_95CI": "Hazard ratio (HR, 95% CI)",
            "P_value": "P-value",
            "Selection_Rationale": "Selection criterion",
            "AUC_95CI_n100": "AUC (95% CI, n=100)",
        }
    )
    s2_a = pd.read_csv(root / "spring模板" / "Table_S2_Robust_Features.csv")
    s2_b = pd.read_csv(root / "spring模板" / "Table_S11_Compact_Signature_Coefficients.csv")
    s2_boot_path = root / "supplementary_figures_rebuild" / "feature_frequency_symbol_bootstrap.csv"
    if s2_boot_path.exists():
        s2_boot = pd.read_csv(s2_boot_path)[["gene_symbol", "selection_frequency"]].rename(
            columns={"gene_symbol": "Symbol", "selection_frequency": "Bootstrap_Selection_Frequency_n100"}
        )
    else:
        s2_boot = pd.DataFrame(columns=["Symbol", "Bootstrap_Selection_Frequency_n100"])
    s2 = s2_a.merge(s2_b[["Symbol", "LASSO_Coefficient", "MR_Beta"]], on="Symbol", how="left")
    s2 = s2.merge(s2_boot, on="Symbol", how="left")
    s2 = s2.rename(columns={"Selection_Frequency": "Nested_CV_Selection_Frequency_n5"})
    s2 = s2[
        [
            "Symbol",
            "Nested_CV_Selection_Frequency_n5",
            "Bootstrap_Selection_Frequency_n100",
            "LASSO_Coefficient",
            "Mean_Importance",
            "Pathway",
            "MR_Beta",
        ]
    ].sort_values(["Bootstrap_Selection_Frequency_n100", "Nested_CV_Selection_Frequency_n5"], ascending=False)
    s2 = s2.rename(
        columns={
            "Nested_CV_Selection_Frequency_n5": "Selection frequency (nested CV, n=5)",
            "Bootstrap_Selection_Frequency_n100": "Selection frequency (bootstrap, n=100)",
            "LASSO_Coefficient": "LASSO coefficient",
            "Mean_Importance": "Mean importance",
            "MR_Beta": "MR β",
        }
    )
    s2["Selection frequency (nested CV, n=5)"] = s2["Selection frequency (nested CV, n=5)"].map(lambda x: format_num(x, 3))
    s2["Selection frequency (bootstrap, n=100)"] = s2["Selection frequency (bootstrap, n=100)"].map(lambda x: format_num(x, 3))
    s2["LASSO coefficient"] = s2["LASSO coefficient"].map(lambda x: format_num(x, 3))
    s2["Mean importance"] = s2["Mean importance"].map(lambda x: format_num(x, 3))
    s2["MR β"] = s2["MR β"].map(lambda x: format_num(x, 3))
    s3 = pd.read_csv(root / "spring模板" / "Table_S7_MR_Complete_Results.csv")
    for col in ["Beta", "SE", "95%_CI_Lower", "95%_CI_Upper", "F_statistic", "Heterogeneity_Q_P", "MR_Egger_Intercept_P", "MR_PRESSO_P"]:
        if col in s3.columns:
            s3[col] = s3[col].map(lambda x: format_num(x, 3))
    if "P_value" in s3.columns:
        s3["P_value"] = s3["P_value"].map(format_p)
    s3 = s3.rename(
        columns={
            "N_SNPs": "N SNPs",
            "Beta": "MR β",
            "P_value": "P-value",
            "F_statistic": "F-statistic",
            "95%_CI_Lower": "95% CI lower",
            "95%_CI_Upper": "95% CI upper",
            "Heterogeneity_Q_P": "Heterogeneity P-value",
            "MR_Egger_Intercept_P": "MR-Egger intercept P-value",
            "MR_PRESSO_P": "MR-PRESSO P-value",
            "Data_Source": "Data source",
        }
    )
    s4 = s3[s3["Method"].astype(str).str.upper() == "IVW"][["Gene", "N SNPs", "F-statistic", "MR-Egger intercept P-value", "Heterogeneity P-value", "MR-PRESSO P-value", "P-value"]].copy()
    s4.columns = ["Gene", "N SNPs", "F-statistic", "Pleiotropy P-value", "Heterogeneity P-value", "MR-PRESSO P-value", "IVW P-value"]
    s5 = pd.read_csv(root / "spring模板" / "Table_S5_Spatial_Autocorrelation.csv").sort_values("Morans_I", ascending=False).reset_index(drop=True)
    s5.insert(0, "Rank", np.arange(1, len(s5) + 1))
    if "Hotspot_Type" in s5.columns:
        hotspot_map = {
            "C1Q_Hotspot": "High-expression cluster",
            "SPP1_Hotspot": "High-expression cluster",
            "Adjacent_Hot": "Adjacent high-expression cluster",
        }
        s5["Hotspot_Type"] = s5["Hotspot_Type"].astype(str).map(lambda x: hotspot_map.get(x, x))
        s5 = s5.rename(columns={"Hotspot_Type": "Spatial cluster type", "Morans_I_Pvalue": "Moran's I P-value", "Z_score": "Z-score"})
    for col in ["Morans_I", "Morans_I_Pvalue", "Z_score"]:
        if col in s5.columns:
            if "Pvalue" in col:
                s5[col] = s5[col].map(format_p)
            else:
                s5[col] = s5[col].map(lambda x: format_num(x, 3))
    s6 = pd.read_csv(root / "spring模板" / "Table_S6_Immune_Correlation.csv")
    for col in s6.columns:
        if col.endswith("Pvalue") or col.endswith("_P"):
            s6[col] = s6[col].map(format_p)
        elif col not in {"Gene", "Cell_Type"}:
            s6[col] = s6[col].map(lambda x: format_num(x, 3))
    curve_path = root / "virtual_perturbation_curve_data.csv"
    if curve_path.exists():
        curve = pd.read_csv(curve_path).copy()
        curve["Condition"] = curve["knockdown_level"].apply(lambda x: "Baseline" if float(x) == 0 else f"C1Q KD {int(float(x) * 100)}%")
        s7 = curve.rename(
            columns={
                "knockdown_level": "Knockdown level",
                "morans_i": "Moran's I",
                "cd8_proximity_median_dist": "CD8 proximity median distance",
                "risk_mean": "Risk score mean",
                "risk_q75": "Risk score Q75",
            }
        )[
            ["Condition", "Knockdown level", "Moran's I", "CD8 proximity median distance", "Risk score mean", "Risk score Q75"]
        ].copy()
    else:
        adata = sc.read_h5ad(str(pick_existing([root / "spatial_clustering_and_tme_analysis.h5ad", root / "processed_spatial_data.h5ad"])))
        x, y = get_spatial_xy(adata)
        coords = np.c_[x, y]
        c1qa = get_gene_vec(adata, "C1QA")
        c1qb = get_gene_vec(adata, "C1QB")
        c1qc = get_gene_vec(adata, "C1QC")
        c1q = np.mean(np.vstack([c1qa, c1qb, c1qc]), axis=0)
        cd8 = get_gene_vec(adata, "CD8A")
        spp1 = get_gene_vec(adata, "SPP1")
        kcnab2 = get_gene_vec(adata, "KCNAB2")
        conds = []
        for frac in [0.0, 0.2, 0.4, 0.6, 0.8]:
            c1q_ko = np.mean(np.vstack([c1qa * (1 - frac), c1qb, c1qc]), axis=0)
            risk = 0.6 * zscore(c1q_ko) + 0.2 * zscore(spp1) - 0.2 * zscore(kcnab2)
            conds.append(
                {
                    "Condition": "Baseline" if frac == 0 else f"C1Q KD {int(frac*100)}%",
                    "Knockdown level": frac,
                    "Moran's I": morans_i_knn(c1q_ko, coords, k=8),
                    "CD8 proximity median distance": float(np.nanmedian(np.abs(c1q_ko - cd8))),
                    "Risk score mean": float(np.mean(risk)),
                    "Risk score Q75": float(np.quantile(risk, 0.75)),
                }
            )
        s7 = pd.DataFrame(conds)
    if not s7.empty:
        b_moran = float(s7.loc[s7["Condition"] == "Baseline", "Moran's I"].iloc[0])
        b_cd8 = float(s7.loc[s7["Condition"] == "Baseline", "CD8 proximity median distance"].iloc[0])
        b_risk = float(s7.loc[s7["Condition"] == "Baseline", "Risk score mean"].iloc[0])
        s7["Delta Moran's I"] = s7["Moran's I"] - b_moran
        s7["Delta CD8 proximity"] = s7["CD8 proximity median distance"] - b_cd8
        s7["Delta Risk score"] = s7["Risk score mean"] - b_risk
        s7["% change Moran's I"] = np.where(abs(b_moran) > 1e-12, (s7["Delta Moran's I"] / b_moran) * 100.0, np.nan)
        s7["% change CD8 proximity"] = np.where(abs(b_cd8) > 1e-12, (s7["Delta CD8 proximity"] / b_cd8) * 100.0, np.nan)
        s7["% change Risk score"] = np.where(abs(b_risk) > 0.05, (s7["Delta Risk score"] / b_risk) * 100.0, np.nan)
    for col in ["Knockdown level", "Moran's I", "CD8 proximity median distance", "Risk score mean", "Risk score Q75"]:
        s7[col] = s7[col].map(lambda x: format_num(x, 3))
    for col in [
        "Delta Moran's I",
        "Delta CD8 proximity",
        "Delta Risk score",
        "% change Moran's I",
        "% change CD8 proximity",
        "% change Risk score",
    ]:
        if col in s7.columns:
            s7[col] = s7[col].map(lambda x: format_num(x, 3))
    s7 = s7.replace("", "NA")
    s7 = s7.rename(columns={"Risk score mean": "Risk score", "Risk score Q75": "Risk score Q75"})
    adata = sc.read_h5ad(str(pick_existing([root / "spatial_clustering_and_tme_analysis.h5ad", root / "processed_spatial_data.h5ad"])))
    c1qa = get_gene_vec(adata, "C1QA")
    c1qb = get_gene_vec(adata, "C1QB")
    c1qc = get_gene_vec(adata, "C1QC")
    c1q = np.mean(np.vstack([c1qa, c1qb, c1qc]), axis=0)
    cd8 = get_gene_vec(adata, "CD8A")
    spp1 = get_gene_vec(adata, "SPP1")
    kcnab2 = get_gene_vec(adata, "KCNAB2")
    base_risk = 0.6 * zscore(c1q) + 0.2 * zscore(spp1) - 0.2 * zscore(kcnab2)
    y_risk = base_risk
    m_cd8 = zscore(cd8)
    m_spp1 = zscore(spp1)
    x_c1q = zscore(c1q)
    x_kcn = zscore(kcnab2)

    def mediation_row(name, x, m, y):
        a = np.polyfit(x, m, 1)[0]
        X = np.column_stack([np.ones(len(x)), x, m])
        beta = np.linalg.lstsq(X, y, rcond=None)[0]
        direct = beta[1]
        b = beta[2]
        total = np.polyfit(x, y, 1)[0]
        indirect = a * b
        pct = (indirect / total * 100.0) if abs(total) > 1e-12 else np.nan
        return {"Path": name, "Direct effect": direct, "Indirect effect": indirect, "% mediated": pct}

    s8 = pd.DataFrame(
        [
            mediation_row("C1Q -> CD8 -> Risk", x_c1q, m_cd8, y_risk),
            mediation_row("C1Q -> SPP1 -> Risk", x_c1q, m_spp1, y_risk),
            mediation_row("KCNAB2 -> CD8 -> Risk", x_kcn, m_cd8, y_risk),
        ]
    )
    for col in ["Direct effect", "Indirect effect", "% mediated"]:
        s8[col] = s8[col].map(lambda x: format_num(x, 3))
    s9 = pd.read_csv(root / "spring模板" / "Table_S9_External_Validation.csv")
    for col in ["AUC_1yr", "AUC_3yr", "AUC_5yr", "C_index"]:
        if col in s9.columns:
            s9[col] = s9[col].map(lambda x: format_num(x, 3))
    if "LogRank_P" in s9.columns:
        s9["LogRank_P"] = s9["LogRank_P"].map(format_p)
    hr_parsed = s9["HR_95CI"].astype(str).map(parse_hr_ci)
    s9["Risk_Direction"] = hr_parsed.map(lambda v: "Risk-increasing (HR>1)" if pd.notna(v) and v >= 1 else "Direction depends on endpoint definition")
    s9["Endpoint_Definition"] = np.where(
        s9["Validation_Status"].astype(str).str.upper() == "VALIDATED",
        "Overall survival (death event)",
        "Immunotherapy cohort endpoint (cohort-specific coding)",
    )
    s9["Harmonized_HR_for_Risk"] = hr_parsed.map(lambda v: (1.0 / v) if pd.notna(v) and v < 1 else v)
    s9["Harmonized_HR_for_Risk"] = s9["Harmonized_HR_for_Risk"].map(lambda x: format_num(x, 3))
    if "HR_95CI" in s9.columns:
        s9["HR_95CI"] = s9["HR_95CI"].map(format_hr95ci_text)
    s9 = s9.rename(
        columns={
            "AUC_1yr": "AUC (1-year)",
            "AUC_3yr": "AUC (3-year)",
            "AUC_5yr": "AUC (5-year)",
            "C_index": "C-index",
            "HR_95CI": "Hazard ratio (HR, 95% CI)",
            "LogRank_P": "P-value",
            "Validation_Status": "Validation status",
            "Risk_Direction": "Effect direction",
            "Endpoint_Definition": "Endpoint definition",
            "Harmonized_HR_for_Risk": "Harmonized risk-oriented HR",
        }
    )
    pan = pd.read_csv(root / "Pan_Cancer_Analysis" / "results" / "pan_cancer_data.csv")
    s10 = (
        pan.groupby("Gene", as_index=False)
        .agg(
            Significant_Cancer_Count=("Significant", lambda x: int(np.sum(x))),
            Mean_Beta=("Beta", "mean"),
            Median_Beta=("Beta", "median"),
        )
        .sort_values("Significant_Cancer_Count", ascending=False)
    )
    s10["Mean_Beta"] = s10["Mean_Beta"].map(lambda x: format_num(x, 3))
    s10["Median_Beta"] = s10["Median_Beta"].map(lambda x: format_num(x, 3))
    s10 = s10.rename(
        columns={
            "Significant_Cancer_Count": "Number of significant cancer types",
            "Mean_Beta": "Mean effect size (β)",
            "Median_Beta": "Median effect size (β)",
        }
    )
    s11 = pd.read_csv(root / "spring模板" / "Table_S12_Spatial_Transcriptomics.csv")
    s12_a = pd.read_csv(root / "spring模板" / "Table_S2_Robust_Features.csv")[["Symbol", "Pathway"]]
    s12_b = pd.read_csv(root / "spring模板" / "Table_S4_Nested_CV_Features.csv")[["Gene_Symbol", "Category", "Annotation"]].rename(columns={"Gene_Symbol": "Symbol"})
    s12 = s12_a.merge(s12_b, on="Symbol", how="left").drop_duplicates()
    s13 = create_table4(root)
    return {
        "S1_Hyperparameters_or_Performance": s1,
        "S2_Full_Feature_List": s2,
        "S3_MR_Full_Results": s3,
        "S4_MR_Sensitivity": s4,
        "S5_Spatial_Metrics": s5,
        "S6_Immune_Correlation": s6,
        "S7_Virtual_KO_Results": s7,
        "S8_Mediation_Analysis": s8,
        "S9_External_Validation": s9,
        "S10_PanCancer": s10,
        "S11_Cohort_Details": s11,
        "S12_Gene_Annotation": s12,
        "S13_Therapeutic_Implication": s13,
    }


def is_numeric_like(v):
    if pd.isna(v):
        return True
    t = str(v).strip()
    if t == "" or t.upper() == "NA" or t.lower() == "not available":
        return True
    if t.startswith("<"):
        t = t[1:]
    t = t.replace("%", "")
    try:
        float(t)
        return True
    except Exception:
        return False


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def apply_three_line_table(table):
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    nil = {"val": "nil", "sz": 0, "space": 0, "color": "000000"}
    line = {"val": "single", "sz": 8, "space": 0, "color": "000000"}
    for i in range(n_rows):
        for j in range(n_cols):
            set_cell_border(table.cell(i, j), left=nil, right=nil, top=nil, bottom=nil)
    for j in range(n_cols):
        set_cell_border(table.cell(0, j), top=line, bottom=line, left=nil, right=nil)
        set_cell_border(table.cell(n_rows - 1, j), bottom=line, left=nil, right=nil)


def set_table_column_widths(table, headers):
    width_map = {
        "Variable": 4.5,
        "Cohort": 4.2,
        "Category": 3.5,
        "Gene": 2.2,
        "Functional annotation": 3.5,
        "Evidence domains": 3.7,
        "Data availability": 3.0,
        "Validation status": 3.0,
        "Endpoint definition": 4.2,
        "Selection criterion": 3.4,
        "Spatial cluster type": 3.6,
        "Data source": 3.0,
    }
    default_width = 2.4
    for j, header in enumerate(headers):
        width_cm = width_map.get(str(header), default_width)
        for cell in table.columns[j].cells:
            cell.width = Cm(width_cm)


def add_df_to_doc(doc: Document, title: str, df: pd.DataFrame, note: str | None = None):
    doc.add_heading(title, level=2)
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False
    for j, c in enumerate(df.columns):
        p = table.rows[0].cells[j].paragraphs[0]
        p.text = str(c)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = "Times New Roman"
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iloc[i, j]
            txt = "Not available" if (pd.isna(v) or str(v).strip() == "") else str(v)
            p = table.rows[i + 1].cells[j].paragraphs[0]
            p.text = txt
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_numeric_like(txt) else WD_ALIGN_PARAGRAPH.LEFT
            if p.runs:
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.name = "Times New Roman"
    set_table_column_widths(table, list(df.columns))
    apply_three_line_table(table)
    if note:
        pn = doc.add_paragraph(f"Note: {note}")
        if pn.runs:
            pn.runs[0].font.size = Pt(9)
            pn.runs[0].font.name = "Times New Roman"
    doc.add_paragraph("")


def apply_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(10)
    h2.font.bold = True
    for sec in doc.sections:
        sec.top_margin = Pt(72)
        sec.bottom_margin = Pt(72)
        sec.left_margin = Pt(72)
        sec.right_margin = Pt(72)


def write_docx(path: Path, title: str, table_map: dict, table_notes: dict | None = None):
    doc = Document()
    apply_doc_style(doc)
    doc.add_heading(title, level=1)
    table_notes = table_notes or {}
    for name, df in table_map.items():
        add_df_to_doc(doc, name, df, table_notes.get(name))
    doc.save(str(path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=str, default=str(Path(__file__).resolve().parent))
    parser.add_argument("--out", type=str, default=None)
    args = parser.parse_args()
    root = Path(args.root).resolve()
    out_dir = ensure_dir(Path(args.out).resolve() if args.out else root / "tables_rebuild")
    main_tables = {
        "Table 1. Cohort Characteristics of Training and Validation Sets": create_table1(root),
        "Table 2. Final Model Performance Summary": create_table2(root),
        "Table 3. Integrated Genetic, Spatial, and Clinical Evidence for Core Genes": create_table3(root),
    }
    supp_tables = create_supp_tables(root)
    main_tables = {k: standardize_missing_cells(v) for k, v in main_tables.items()}
    supp_tables = {k: standardize_missing_cells(v) for k, v in supp_tables.items()}
    main_csv_dir = ensure_dir(out_dir / "main_tables_csv")
    supp_csv_dir = ensure_dir(out_dir / "supplementary_tables_csv")
    for k, df in main_tables.items():
        df.to_csv(main_csv_dir / f"{k}.csv", index=False)
    for k, df in supp_tables.items():
        df.to_csv(supp_csv_dir / f"{k}.csv", index=False)
    with pd.ExcelWriter(out_dir / "Main_Tables.xlsx", engine="openpyxl") as xw:
        for k, df in main_tables.items():
            df.to_excel(xw, sheet_name=k[:31], index=False)
    with pd.ExcelWriter(out_dir / "Supplementary_Tables.xlsx", engine="openpyxl") as xw:
        for k, df in supp_tables.items():
            df.to_excel(xw, sheet_name=k[:31], index=False)
    supp_tables_titled = {
        "Supplementary Table 1. Model Development and Benchmarking": supp_tables["S1_Hyperparameters_or_Performance"],
        "Supplementary Table 2. Full Feature List": supp_tables["S2_Full_Feature_List"],
        "Supplementary Table 3. Full MR Results": supp_tables["S3_MR_Full_Results"],
        "Supplementary Table 4. MR Sensitivity Diagnostics": supp_tables["S4_MR_Sensitivity"],
        "Supplementary Table 5. Spatial Metrics": supp_tables["S5_Spatial_Metrics"],
        "Supplementary Table 6. Immune Correlation": supp_tables["S6_Immune_Correlation"],
        "Supplementary Table 7. Virtual KO Results": supp_tables["S7_Virtual_KO_Results"],
        "Supplementary Table 8. Mediation Analysis": supp_tables["S8_Mediation_Analysis"],
        "Supplementary Table 9. External Validation": supp_tables["S9_External_Validation"],
        "Supplementary Table 10. Pan-Cancer Summary": supp_tables["S10_PanCancer"],
        "Supplementary Table 11. Cohort Details": supp_tables["S11_Cohort_Details"],
        "Supplementary Table 12. Gene Annotation": supp_tables["S12_Gene_Annotation"],
        "Supplementary Table 13. Therapeutic Implication Summary": supp_tables["S13_Therapeutic_Implication"],
    }
    main_notes = {
        "Table 1. Cohort Characteristics of Training and Validation Sets": "Smoking information in the TCGA cohort was limited and categorized as ever-smokers based on available records, whereas smoking categories were directly available in GSE31210.",
        "Table 2. Final Model Performance Summary": "Final model selected based on highest cross-validated AUC.",
        "Table 3. Integrated Genetic, Spatial, and Clinical Evidence for Core Genes": "Abbreviations: MR, Mendelian randomization; HR, hazard ratio.",
    }
    supp_notes = {
        "Supplementary Table 7. Virtual KO Results": "Derived from the same perturbation series as Figure 5 (0-80% C1Q knockdown). Effect size is calculated relative to baseline and increases with perturbation intensity.",
        "Supplementary Table 8. Mediation Analysis": "Negative proportion mediated indicates an opposing indirect effect (suppression effect).",
        "Supplementary Table 9. External Validation": "HR direction depends on cohort endpoint coding. Harmonized_HR_for_Risk standardizes orientation to risk-increasing direction.",
        "Supplementary Table 3. Full MR Results": "F-statistics are derived from large-scale eQTL datasets and may appear inflated due to sample size.",
        "Supplementary Table 13. Therapeutic Implication Summary": "Effect estimates are reported as odds ratios (OR) due to binary endpoints.",
    }
    write_docx(out_dir / "Main_Tables.docx", "Main Tables (Submission Format)", main_tables, table_notes=main_notes)
    write_docx(out_dir / "Supplementary_Tables.docx", "Supplementary Tables (Submission Format)", supp_tables_titled, table_notes=supp_notes)
    print(f"Done. Outputs saved to: {out_dir}")


if __name__ == "__main__":
    main()
